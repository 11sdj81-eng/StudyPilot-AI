#!/usr/bin/env python3
"""Recover TRUE_FINAL layout — content IDENTICAL, only font/spacing/photo adjusted.

Baseline: FINAL_AI_RESUME_TRUE_FINAL.docx
Content: ZERO changes. Only visual adjustments.
"""

import subprocess, sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

# ═══════════════════ Content — VERBATIM from TRUE_FINAL ═══════════════════

NAME = "苏东健"
PHONE = "15299730320"
EMAIL = "529412512@qq.com"
GITHUB_URL = "github.com/11sdj81-eng"
DIRECTION = "AI 应用开发实习生"

EDUCATION = dict(
    school="北京邮电大学", major="电子信息工程", degree="本科大二（2024.09 - 至今）",
    courses="程序设计基础、数据结构、概率论与随机过程、信号与系统、数字电路与逻辑设计",
)

PROJECTS = [
    {
        "name": "StudyPilot AI ｜ AI 学习教练系统",
        "time": "2026.04 - 至今",
        "desc": "面向大学生的 AI 学习教练系统。覆盖 3 门课程、80+ 教材 Chunk、45 条 Benchmark Query，输出 14 份系统评测与审计报告。",
        "tech": "Python  |  Streamlit  |  DeepSeek API  |  FAISS  |  BM25  |  RRF  |  pytest",
        "work": [
            "设计课程级 AI Tutor 架构，统一管理意图识别、课程资源调度、RAG 检索增强与 PDF 生成流程，形成「对话→检索→生成→交付」完整教学闭环。",
            "自研 Hybrid Retrieval 混合检索引擎：FAISS 语义检索（BGE Embedding）+ BM25 字符级关键词检索（Bigram/Trigram 分词，标准 BM25 公式，argpartition O(n) Top-K）+ RRF 倒数秩融合。45 条跨课程 Benchmark 测试：Top1 Accuracy 从纯 FAISS 的 46.67% 提升至 Hybrid 的 63.33%（+35.7%），Top3 Recall 从 50.00% 提升至 66.67%（+33.3%），Citation Accuracy 达到 100%。",
            "设计 Citation Quality Gate 引用质量过滤器：基于汉字占比 / 符号噪声率 / 目录特征 / 主题匹配四维评分，自动拒绝目录页、OCR 乱码页与低分 chunks（平均拒绝率 94.67%），确保 LLM Prompt 中每一条教材引用均可信且可溯源。",
            "构建 Mastery Score（掌握度评分）与 Wrong Question Memory（错题记忆）系统：correct/wrong 计数 → mastery 0-100 评分（未测试=50 中性分），驱动出题与模拟卷优先覆盖薄弱知识点，形成「答题→评测→反馈→针对性出题」的个性化学习闭环。",
        ],
    },
    {
        "name": "SmartVoiceSystem ｜ 多模态智能家居控制系统",
        "time": "2026.03 - 2026.05",
        "desc": "基于 Python 的智能家居 AI 控制平台，支持文本、语音、手势三种输入通道，接入 DeepSeek 进行自然语言意图解析，统一通过 JSON Action Schema 映射为设备控制指令，实现灯光、空调、风扇的跨模态统一控制。",
        "tech": "Python  |  DeepSeek API  |  OpenCV  |  MediaPipe  |  Tkinter  |  SpeechRecognition",
        "work": [
            "设计多模态输入架构，文本（Tkinter 聊天面板）、语音（SpeechRecognition 模块）、手势（OpenCV + MediaPipe 21 点手部关键点提取）三种通道并行处理，通过 JSON Action Schema 统一映射为设备动作指令，实现灯光开关/亮度、空调模式/温度、风扇档位的跨模态控制。",
            "接入 DeepSeek API 进行自然语言意图解析，将「帮我设置舒服的学习环境」等口语输入通过 Prompt Engineering（提示词工程）解析为结构化设备动作序列（灯光→调暗、空调→26°C→低风、风扇→低速），支持学习/睡眠/离家三种场景模式自动切换与多设备联动。",
            "构建实时手势识别管道：基于人手 21 关键点坐标 + 距离/角度特征工程，实现手掌张开（全开灯）、握拳（全关）、双指滑动（调亮度）等 5+ 种手势分类，识别帧率 25-30fps，结合 Tkinter GUI 实现设备状态实时反馈。",
        ],
    },
]

SKILLS = [
    dict(cat="AI 应用开发", text="DeepSeek API、RAG、Prompt Engineering（提示词工程）、OpenAI API（了解）"),
    dict(cat="检索系统", text="FAISS、BM25、RRF、Hybrid Retrieval Benchmark"),
    dict(cat="编程语言", text="Python 项目开发、C 语言基础、数据结构与算法基础、LeetCode 50+ 题"),
    dict(cat="工程实践", text="REST API、JSON 数据处理、Git（基础使用）、pytest（基础测试）"),
]

HONORS = [
    "校级创新创业项目二等奖（数学建模方向）",
    "小班学习委员（2025.09 - 2026.01）",
    "校志愿者协会成员（2024.09 - 2025.06）",
    "CET-6（436）",
]

# ═══════════════════ Typst Template — Layout Only ═══════════════════

TEMPLATE = r"""// Layout Recovered — TRUE_FINAL content, optimized readability

#let recovered(
  name: "", phone: "", email: "", github_url: "", direction: "",
  education: (:), projects: (), skills: (), honors: ()
) = {
  set page(paper: "a4", margin: (top: 0.25cm, bottom: 0.0cm, left: 1.3cm, right: 1.3cm))
  set text(font: ("Heiti SC", "PingFang SC"), size: 10pt, lang: "zh")
  set par(leading: 0.28em, first-line-indent: 0pt)

  let dark-blue = rgb(0, 51, 102)
  let sep = line(length: 100%, stroke: 0.4pt + rgb(180, 180, 180))
  let gray-sep = line(length: 100%, stroke: 0.3pt + rgb(210, 210, 210))

  // ═══ Header ═══
  grid(columns: (1fr, 2.3cm), column-gutter: 0.5cm,
    {
      text(size: 18pt, weight: "bold")[#name]
      v(0.02em)
      text(size: 10pt, fill: rgb(90, 90, 90))[#phone  |  #email]
      v(0.01em)
      text(size: 10pt, weight: "bold")[*GitHub：*#github_url]
      v(0.01em)
      text(size: 10pt)[*开源项目：*StudyPilot-AI、SmartVoiceSystem]
      v(0.01em)
      text(size: 10pt, fill: rgb(70, 70, 70))[*求职意向：*#direction]
    },
    {
      rect(width: 2.0cm, height: 2.8cm, stroke: 0.3pt + rgb(160, 160, 160), fill: rgb(250, 250, 250))[
        #set text(size: 7pt, fill: gray)
        #align(center + horizon)[证件照]
      ]
    },
  )

  v(0.06em)
  line(length: 100%, stroke: 0.5pt + black)

  // ═══ Education ═══
  v(0.02em)
  text(size: 11pt, weight: "bold", fill: dark-blue)[教育背景]
  v(0.01em)
  text(size: 10pt)[*#education.school*  |  #education.major  |  #education.degree]
  if "courses" in education {
    v(0.01em)
    text(size: 10pt)[主修课程：#education.courses]
  }

  // ═══ Projects ═══
  v(0.06em)
  sep
  v(0.02em)
  text(size: 11pt, weight: "bold", fill: dark-blue)[项目经历]

  // ── Project 1 ──
  v(0.04em)
  grid(columns: (1fr, auto), column-gutter: 0.3cm,
    { text(size: 10.5pt, weight: "bold")[#projects.at(0).name] },
    { text(size: 10pt, fill: rgb(100, 100, 100))[#projects.at(0).time] },
  )
  v(0.01em)
  text(size: 10pt)[*项目简介：*#projects.at(0).desc]
  v(0.01em)
  text(size: 10pt)[*技术栈：*#projects.at(0).tech]
  v(0.01em)
  text(size: 10pt)[*主要工作：*]
  v(0.01em)
  for bullet in projects.at(0).work {
    text(size: 10pt)[- #bullet]
    v(0.01em)
  }

  // ── Separator ──
  v(0.02em)
  gray-sep
  v(0.02em)

  // ── Project 2 ──
  grid(columns: (1fr, auto), column-gutter: 0.3cm,
    { text(size: 10.5pt, weight: "bold")[#projects.at(1).name] },
    { text(size: 10pt, fill: rgb(100, 100, 100))[#projects.at(1).time] },
  )
  v(0.01em)
  text(size: 10pt)[*项目简介：*#projects.at(1).desc]
  v(0.01em)
  text(size: 10pt)[*技术栈：*#projects.at(1).tech]
  v(0.01em)
  text(size: 10pt)[*主要工作：*]
  v(0.01em)
  for bullet in projects.at(1).work {
    text(size: 10pt)[- #bullet]
    v(0.01em)
  }

  // ═══ Skills ═══
  v(0.02em)
  sep
  v(0.02em)
  text(size: 11pt, weight: "bold", fill: dark-blue)[技术能力]
  v(0.03em)
  for skill in skills {
    text(size: 10pt)[*#skill.cat：*#skill.text]
    v(0.03em)
  }

  // ═══ Honors ═══
  v(0.02em)
  sep
  v(0.02em)
  text(size: 11pt, weight: "bold", fill: dark-blue)[荣誉与校园经历]
  v(0.03em)
  for h in honors {
    text(size: 10pt)[- #h]
    v(0.03em)
  }
}
"""

# ═══════════════════ Bold Key Terms ═══════════════════

BOLD_TERMS = [
    # StudyPilot
    "Hybrid Retrieval", "FAISS", "BM25", "RRF", "BGE Embedding",
    "Top1 Accuracy", "46.67%", "63.33%", "35.7%",
    "Top3 Recall", "50.00%", "66.67%", "33.3%",
    "Citation Accuracy", "100%",
    "Citation Quality Gate", "94.67%",
    "Mastery Score", "Wrong Question Memory",
    # SmartVoiceSystem
    "DeepSeek API", "JSON Action Schema", "Prompt Engineering",
    "OpenCV", "MediaPipe", "25-30fps",
]

import re as _re

def bold_terms_in(text: str) -> str:
    """Wrap key terms in Typst bold markers *...*."""
    result = text
    for term in sorted(BOLD_TERMS, key=len, reverse=True):
        escaped = _re.escape(term)
        result = _re.sub(rf'(?<!\*){escaped}(?!\*)', f'*{term}*', result)
    return result

# ═══════════════════ Typst Generator ═══════════════════

def esc(s):
    return s.replace("→", "->").replace("←", "<-").replace("“", "「").replace("”", "」").replace('"', "'")

def generate_typst():
    lines = [
        '#import "templates/layout_recovered.typ": recovered',
        "",
        "#show: doc => recovered(",
        f'  name: "{NAME}", phone: "{PHONE}", email: "{EMAIL}", github_url: "{GITHUB_URL}", direction: "{DIRECTION}",',
        f'  education: (school: "{EDUCATION["school"]}", major: "{EDUCATION["major"]}", degree: "{EDUCATION["degree"]}", courses: "{EDUCATION["courses"]}"),',
        "", "  projects: (",
    ]
    for pi, proj in enumerate(PROJECTS):
        lines.append("    (")
        lines.append(f'      name: "{esc(proj["name"])}", time: "{esc(proj["time"])}",')
        lines.append(f'      desc: "{esc(proj["desc"])}",')
        lines.append(f'      tech: "{esc(proj["tech"])}",')
        lines.append("      work: (")
        for bi, b in enumerate(proj["work"]):
            lines.append(f'        "{esc(bold_terms_in(b))}"{"," if bi < len(proj["work"])-1 else ""}')
        lines.append("      )")
        lines.append(f"    ){',' if pi < len(PROJECTS)-1 else ''}")
    lines.append("  ),")
    lines.append("")
    lines.append("  skills: (")
    for i, s in enumerate(SKILLS):
        lines.append(f'    (cat: "{esc(s["cat"])}", text: "{esc(s["text"])}"){"," if i < len(SKILLS)-1 else ""}')
    lines.append("  ),")
    lines.append("")
    lines.append("  honors: (")
    for i, h in enumerate(HONORS):
        lines.append(f'    "{esc(h)}"{"," if i < len(HONORS)-1 else ""}')
    lines.append("  )")
    lines.append(")")
    return "\n".join(lines) + "\n"

# ═══════════════════ DOCX Generator ═══════════════════

def generate_docx(path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin, sec.bottom_margin = Cm(0.6), Cm(0.4)
    sec.left_margin, sec.right_margin = Cm(1.5), Cm(1.5)
    s = doc.styles['Normal']
    s.font.size = Pt(10)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.line_spacing = 1.15

    DARK_BLUE = (0, 51, 102)
    GRAY = (100, 100, 100)

    def L(text="", bold=False, size=10, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        if text:
            r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
            r.font.name = 'Arial'
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts', {})
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Microsoft YaHei')
            rPr.insert(0, rFonts)
            if color: r.font.color.rgb = RGBColor(*color)
        return p

    def Sep(): L("─"*72, size=5, color=(180, 180, 180))
    def GraySep(): L("─"*72, size=5, color=(210, 210, 210))

    L(NAME, bold=True, size=20)
    L(f"{PHONE}  |  {EMAIL}", size=10, color=(90, 90, 90))
    L(f"GitHub：{GITHUB_URL}", bold=True, size=10)
    L("开源项目：StudyPilot-AI、SmartVoiceSystem", size=10)
    L(f"求职意向：{DIRECTION}", size=10, color=(70, 70, 70))
    L("─"*72, size=5, color=(0, 0, 0))

    L("教育背景", bold=True, size=12, color=DARK_BLUE)
    L(f"{EDUCATION['school']}  |  {EDUCATION['major']}  |  {EDUCATION['degree']}", size=10)
    L(f"主修课程：{EDUCATION['courses']}", size=10)
    Sep()

    L("项目经历", bold=True, size=12, color=DARK_BLUE)
    for pi, proj in enumerate(PROJECTS):
        if pi == 1:
            GraySep()  # Separator between projects
        L(f"{proj['name']}  |  {proj['time']}", bold=True, size=11)
        L(f"项目简介：{proj['desc']}", size=10)
        L(f"技术栈：{proj['tech']}", size=10)
        L("主要工作：", size=10)
        for w in proj['work']:
            L(f"  - {w}", size=10)
    Sep()

    L("技术能力", bold=True, size=12, color=DARK_BLUE)
    for s in SKILLS:
        L(f"{s['cat']}：{s['text']}", size=10)
    Sep()

    L("荣誉与校园经历", bold=True, size=12, color=DARK_BLUE)
    for h in HONORS:
        L(f"  - {h}", size=10)

    doc.save(str(path))

# ═══════════════════ Main ═══════════════════

def main():
    print("=" * 60)
    print("TRUE_FINAL Layout Recovery")
    print("=" * 60)

    # Template
    tmpl_path = ROOT / "templates/layout_recovered.typ"
    tmpl_path.write_text(TEMPLATE, encoding="utf-8")
    print(f"✅ {tmpl_path.name}")

    # Typst → PDF
    typ_path = ROOT / "FINAL_AI_RESUME_LAYOUT_RECOVERED.typ"
    typ_path.write_text(generate_typst(), encoding="utf-8")
    print(f"✅ {typ_path.name}")

    pdf_path = ROOT / "FINAL_AI_RESUME_LAYOUT_RECOVERED.pdf"
    r = subprocess.run(["typst", "compile", str(typ_path), str(pdf_path), "--root", str(ROOT)],
                       capture_output=True, text=True, cwd=str(ROOT))
    if r.returncode != 0:
        print(f"❌ PDF: {r.stderr[:500]}"); return 1
    print(f"✅ {pdf_path.name} ({pdf_path.stat().st_size//1024} KB)")

    # DOCX
    docx_path = ROOT / "FINAL_AI_RESUME_LAYOUT_RECOVERED.docx"
    generate_docx(docx_path)
    print(f"✅ {docx_path.name} ({docx_path.stat().st_size//1024} KB)")

    # PNG
    try:
        import fitz
        doc = fitz.open(str(pdf_path))
        pix = doc[0].get_pixmap(dpi=150)
        png_path = ROOT / "FINAL_AI_RESUME_LAYOUT_RECOVERED.png"
        pix.save(str(png_path))
        print(f"✅ {png_path.name} ({png_path.stat().st_size//1024} KB, {pix.width}×{pix.height})")
    except Exception as e:
        print(f"⚠️ PNG: {e}")

    # Verify
    import fitz
    doc = fitz.open(str(pdf_path))
    t = doc[0].get_text()
    print(f"\nPDF: {len(doc)}p, {len(t)}c, {len(t.splitlines())}l")
    # Verify content matches TRUE_FINAL
    true_text = fitz.open(str(ROOT / "FINAL_AI_RESUME_TRUE_FINAL.pdf"))[0].get_text()
    # Normalize whitespace for comparison
    def norm(s):
        import re
        return re.sub(r'\s+', '', s)
    content_match = norm(t) == norm(true_text)
    print(f"  {'✅' if content_match else '❌'} Content matches TRUE_FINAL: {len(norm(t))} vs {len(norm(true_text))} chars (normalized)")
    if not content_match:
        # Show first diff
        for i, (a, b) in enumerate(zip(norm(t), norm(true_text))):
            if a != b:
                ctx = max(0, i-20)
                print(f"  Diff at pos {i}: ...{norm(t)[ctx:i+20]}... vs ...{norm(true_text)[ctx:i+20]}...")
                break
    print("Done.")

if __name__ == "__main__":
    sys.exit(main())
